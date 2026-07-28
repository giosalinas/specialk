import streamlit as st
import time
import requests
import os
import html
import threading
from datetime import datetime

# --- 1. ROBUST IMPORTS ---
# Checks if libraries exist; if not, disables that feature gracefully
try:
    import pypdf
except ImportError:
    pypdf = None
try:
    import docx
except ImportError:
    docx = None

from audio_engine import AudioRecorder, MLX_MODELS, WHISPERX_MODELS
from meeting_manager import (
    load_meetings, save_meeting, delete_meeting, update_meeting_chat,
    load_prompts, add_prompt, delete_prompt, update_prompt,
    load_settings, save_settings, add_files_to_meeting,
    increment_prompt_usage, get_top_prompts
)

# --- CONFIG ---
st.set_page_config(page_title="specialk", layout="wide", page_icon="🥣")
LLM_URL = "http://localhost:1234/v1/chat/completions"

# --- HELPER: AUTO-SAVE CALLBACK ---
def update_model_setting():
    if "active_llm_widget" in st.session_state:
        new_model = st.session_state.active_llm_widget
        current_settings = load_settings()
        current_settings["llm_model_id"] = new_model
        save_settings(current_settings)

# --- HELPER: FILE PARSER ---
# Reads text from uploaded files to build LLM context
def get_text_from_files(file_list):
    context_str = ""
    errors = []

    for file in file_list:
        fname = file.name.lower()
        try:
            file.seek(0)
            text = ""

            # 1. PDF Handling
            if fname.endswith('.pdf'):
                if pypdf:
                    pdf = pypdf.PdfReader(file)
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted: text += extracted + "\n"
                    context_str += f"\n\n--- FILE: {file.name} (PDF) ---\n{text}"
                else:
                    errors.append(f"❌ Missing 'pypdf' library. Cannot read {file.name}")

            # 2. Word Doc Handling
            elif fname.endswith('.docx'):
                if docx:
                    doc = docx.Document(file)
                    text = "\n".join([p.text for p in doc.paragraphs])
                    context_str += f"\n\n--- FILE: {file.name} (DOCX) ---\n{text}"
                else:
                    errors.append(f"❌ Missing 'python-docx' library. Cannot read {file.name}")

            # 3. Plain Text Handling
            elif fname.endswith(('.txt', '.md', '.csv', '.json', '.py')):
                text = file.read().decode("utf-8")
                context_str += f"\n\n--- FILE: {file.name} ---\n{text}"

            # Validation
            if not text and not errors and fname.endswith(('.pdf', '.docx', '.txt')):
                errors.append(f"⚠️ {file.name} was empty or unreadable (Scanned?).")

        except Exception as e:
            errors.append(f"❌ Error reading {file.name}: {str(e)}")

    return context_str, errors

# --- HELPER: LLM FUNCTION ---
def ask_llm(prompt, transcript_text, files_text, use_history=False):
    settings = load_settings()
    llm_id = settings.get("llm_model_id", "local-model")
    all_prompts = load_prompts()
    is_custom = any(p['text'].strip() == prompt.strip() for p in all_prompts)

    # Load style preferences (with defaults)
    if "style_detail" not in st.session_state: st.session_state.style_detail = "Brief & Concise (Default)"
    if "style_tone" not in st.session_state: st.session_state.style_tone = "Neutral (Default)"
    if "style_lang" not in st.session_state: st.session_state.style_lang = "Human (Default)"

    # Build context with clear separation
    context_block = ""
    if transcript_text.strip():
        context_block += f"\n\n=== MEETING TRANSCRIPT (Primary Content) ===\n{transcript_text}"
    if files_text.strip():
        context_block += f"\n\n=== REFERENCE DOCUMENTS (Supporting Context) ===\n{files_text}"

    base_instructions = """You are a meeting assistant.
IMPORTANT: The MEETING TRANSCRIPT is the primary content - this is what the meeting is about.
The REFERENCE DOCUMENTS are supplementary materials uploaded for context/reference only.
When asked questions, focus on the MEETING TRANSCRIPT not the referenced documents.
Use the REFERENCE DOCUMENTS only to provide additional context, provide better answers or clarify topics discussed in the meeting."""

    if is_custom:
        system_content = f"{base_instructions}{context_block}"
    else:
        detail_instr = STYLE_PROMPTS["detail"][st.session_state.style_detail]
        tone_instr = STYLE_PROMPTS["tone"][st.session_state.style_tone]
        lang_instr = STYLE_PROMPTS["language"][st.session_state.style_lang]
        style_block = f"{detail_instr} {tone_instr} {lang_instr}".strip()
        system_content = f"{base_instructions}\nSTYLE: {style_block}{context_block}"

    # Store the full prompt for debug display
    st.session_state.debug_system_prompt = system_content
    st.session_state.debug_user_prompt = prompt

    messages = [{"role": "system", "content": system_content}]
    if use_history:
        messages.extend(st.session_state.chat_history)
    messages.append({"role": "user", "content": prompt})

    try:
        payload = {"model": llm_id, "messages": messages, "temperature": 0.7, "stream": False}
        start_time = time.perf_counter()
        response = requests.post(LLM_URL, json=payload, timeout=300)
        elapsed = round(time.perf_counter() - start_time, 2)

        if response.status_code == 200:
            content = response.json()['choices'][0]['message']['content']
            return f"{content}\n\n---\n*⏱️ {elapsed}s | Model: `{llm_id}`*"
        else: return f"Error {response.status_code}: {response.text}"
    except Exception as e: return f"Connection Error: {e}"


# --- HELPER: LLM TRANSCRIPTION FORMATTER ---
def format_transcription_with_llm(raw_transcript):
    """
    Uses LLM to clean and format raw Whisper transcription into coherent, readable text.
    Strips timestamps before sending to avoid hallucination.
    """
    settings = load_settings()
    llm_id = settings.get("llm_transcription_model_id", "local-model")

    # Remove timestamps to prevent LLM from hallucinating times
    import re
    clean_transcript = re.sub(r'\[[\d:]+\]\s*', '', raw_transcript)

    system_prompt = """You are a Text Restoration Engine. Your task is mechanical: Restore punctuation and casing to the raw text stream provided.

### INPUT
A raw, unformatted transcript with transcription errors and stuttering.

### OUTPUT
A solid block of text with correct punctuation.

### RESTRICTIONS (STRICT)
1.  **NO SUMMARIZATION:** Do not summarize the text. Do not create headers (e.g., "**Discussion**"). Do not use bullet points.
2.  **WORD FIDELITY:** Do not add words that are not in the source text.
    * *Bad:* "We should focus on personal growth." (Added content)
    * *Good:* "The event support isn't nailed down." (Original content)
3.  **NO REORDERING:** Do not move sentences around. Keep the chronological flow exactly as it appears.
4.  **DELETION ONLY:** You are allowed to DELETE filler words (um, uh, you know, sort of) and stuttering (I... I... I). You are NOT allowed to ADD transition words (Therefore, Furthermore, In conclusion).

### HANDLING FRAGMENTS
If a sentence is incomplete in the raw text, leave it incomplete in your output. Do not finish the thought for the speaker.

### EXAMPLE
* *Raw:* "i think we should uh go to the... to the store and buy"
* *Wrong:* "I think we should go to the store and buy supplies." (Hallucinated "supplies")
* *Correct:* "I think we should go to the store and buy."

### BEGIN PROCESSING
Output ONLY the restored text. Do not output intro sentences or conversational filler."""

    user_content = f"""Clean this transcript. Output ONLY the cleaned text with paragraph breaks between topics:

{clean_transcript}"""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content}
    ]

    try:
        payload = {"model": llm_id, "messages": messages, "temperature": 0.1, "stream": False}
        response = requests.post(LLM_URL, json=payload, timeout=90)

        if response.status_code == 200:
            result = response.json()['choices'][0]['message']['content'].strip()
            if len(result) < 20:
                print(f"LLM output too short ({len(result)}), using raw")
                return None
            return result
        else:
            print(f"LLM request failed with status {response.status_code}")
            return None
    except Exception as e:
        print(f"LLM Transcription Error: {e}")
        return None


class LLMTranscriptionManager:
    """Manages LLM formatting - triggers based on new word count, runs in background thread."""

    def __init__(self):
        self.formatted_output = ""
        self.committed_formatted = ""  # Older text that won't change
        self.is_formatting = False
        self.lock = threading.Lock()
        self.last_word_count = 0  # Track words processed
        self.min_new_words = 40  # Trigger after ~40 new words
        self.last_committed_raw_pos = 0
        self.reprocess_window = 2000
        self.pending_raw_text = None
        self.start_time = None  # Track when recording started
        self.paragraph_count = 0  # Track paragraphs for approximate timestamps

    def reset(self):
        """Reset the manager for a new session."""
        with self.lock:
            self.formatted_output = ""
            self.committed_formatted = ""
            self.is_formatting = False
            self.last_word_count = 0
            self.last_committed_raw_pos = 0
            self.pending_raw_text = None
            self.start_time = time.time()
            self.paragraph_count = 0

    def _get_approx_timestamp(self):
        """Get approximate timestamp based on elapsed time."""
        if self.start_time is None:
            self.start_time = time.time()
        elapsed = int(time.time() - self.start_time)
        hours = elapsed // 3600
        minutes = (elapsed % 3600) // 60
        seconds = elapsed % 60
        return f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

    def _count_words(self, text):
        """Count words in text, ignoring timestamps."""
        import re
        clean = re.sub(r'\[[\d:]+\]', '', text)
        return len(clean.split())

    def should_format(self, current_raw_text):
        """Check if we have enough new words to trigger formatting."""
        current_word_count = self._count_words(current_raw_text)
        new_words = current_word_count - self.last_word_count
        return new_words >= self.min_new_words and not self.is_formatting

    def _find_safe_commit_point(self, raw_text, target_pos):
        """Find a safe position to commit (at a timestamp boundary)."""
        search_start = max(0, target_pos - 200)
        search_end = min(len(raw_text), target_pos + 200)
        search_area = raw_text[search_start:search_end]

        import re
        timestamps = list(re.finditer(r'\n?\[[\d:]+\]', search_area))
        if timestamps:
            for ts in reversed(timestamps):
                abs_pos = search_start + ts.start()
                if abs_pos <= target_pos:
                    return abs_pos
        return target_pos

    def _add_timestamps_to_paragraphs(self, text):
        """Add approximate timestamps to the start of each paragraph."""
        if not text:
            return text

        paragraphs = text.strip().split('\n\n')
        result_paragraphs = []

        for para in paragraphs:
            para = para.strip()
            if para:
                # Add timestamp to paragraph
                ts = self._get_approx_timestamp()
                result_paragraphs.append(f"{ts} {para}")
                self.paragraph_count += 1

        return '\n\n'.join(result_paragraphs)

    def _do_format_in_background(self, raw_text):
        """Background thread function that does the actual LLM formatting."""
        try:
            clean_raw = raw_text.replace(' ▌', '').strip()
            raw_length = len(clean_raw)

            if raw_length > self.reprocess_window:
                commit_target = raw_length - self.reprocess_window
                commit_pos = self._find_safe_commit_point(clean_raw, commit_target)

                if commit_pos > self.last_committed_raw_pos:
                    new_commit_raw = clean_raw[self.last_committed_raw_pos:commit_pos].strip()
                    if new_commit_raw:
                        formatted_chunk = format_transcription_with_llm(new_commit_raw)
                        if formatted_chunk:
                            # Add timestamps to formatted chunk
                            formatted_chunk = self._add_timestamps_to_paragraphs(formatted_chunk)
                            with self.lock:
                                if self.committed_formatted:
                                    self.committed_formatted += "\n\n" + formatted_chunk
                                else:
                                    self.committed_formatted = formatted_chunk
                                self.last_committed_raw_pos = commit_pos

                recent_raw = clean_raw[self.last_committed_raw_pos:].strip()
            else:
                recent_raw = clean_raw

            if recent_raw:
                formatted_recent = format_transcription_with_llm(recent_raw)

                if formatted_recent and formatted_recent.strip():
                    # Add timestamps to formatted output
                    formatted_recent = self._add_timestamps_to_paragraphs(formatted_recent)
                    with self.lock:
                        if self.committed_formatted:
                            self.formatted_output = self.committed_formatted + "\n\n" + formatted_recent
                        else:
                            self.formatted_output = formatted_recent
                        self.last_word_count = self._count_words(raw_text)
                else:
                    with self.lock:
                        if self.committed_formatted:
                            self.formatted_output = self.committed_formatted + "\n\n" + recent_raw
                        else:
                            self.formatted_output = recent_raw
                        self.last_word_count = self._count_words(raw_text)

        except Exception as e:
            print(f"LLM Transcription Manager Error: {e}")
        finally:
            with self.lock:
                self.is_formatting = False

    def update_and_format(self, current_raw_text):
        """
        Non-blocking: triggers formatting in background thread if needed.
        Always returns immediately with current formatted output.
        """
        if not current_raw_text or len(current_raw_text.strip()) < 50:
            return self.formatted_output

        if not self.should_format(current_raw_text):
            return self.formatted_output

        with self.lock:
            if self.is_formatting:
                return self.formatted_output
            self.is_formatting = True

        format_thread = threading.Thread(
            target=self._do_format_in_background,
            args=(current_raw_text,),
            daemon=True
        )
        format_thread.start()

        return self.formatted_output

    def get_formatted_output(self):
        """Get the current formatted output."""
        with self.lock:
            return self.formatted_output
            return self.formatted_output


# --- STATE & SETTINGS ---
@st.cache_resource
def get_recorder(): return AudioRecorder(load_settings())
recorder = get_recorder()

@st.cache_resource
def get_llm_transcription_manager():
    return LLMTranscriptionManager()
llm_transcription_manager = get_llm_transcription_manager()

if "page" not in st.session_state: st.session_state.page = "home"
if "current_meeting_id" not in st.session_state: st.session_state.current_meeting_id = None
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "show_all_prompts" not in st.session_state: st.session_state.show_all_prompts = False
if "temp_title" not in st.session_state: st.session_state.temp_title = "Untitled"
if "temp_files" not in st.session_state: st.session_state.temp_files = []

STYLE_PROMPTS = {
    "detail": {
        "Brief & Concise (Default)": "Provide short, focused responses. Use bullet points when listing multiple items. Limit explanations to 2-3 sentences maximum. Omit unnecessary context or background information.",
        "Brief (One Sentence)": "Respond with exactly one sentence. Be direct and to the point. Do not elaborate or provide additional context.",
        "Detailed": "Provide comprehensive, thorough responses. Include relevant context, examples, and explanations. Break down complex topics into clear sections. Feel free to elaborate on important points."
    },
    "tone": {
        "Neutral (Default)": "Use a professional, balanced tone. Be objective and straightforward without being cold or overly formal.",
        "Friendly": "Use a warm, conversational tone. Be approachable and personable while remaining helpful. Use casual language where appropriate.",
        "Formal": "Use a formal, professional tone suitable for business communication. Avoid contractions and colloquialisms. Maintain a respectful, polished style."
    },
    "language": {
        "Human (Default)": "Write naturally as a human would. Use varied sentence structures, everyday vocabulary, and a conversational flow. Avoid robotic or overly structured phrasing.",
        "AI Normal": "You may use structured formatting, technical terminology, and systematic organization. Clarity and precision take priority over conversational style."
    }
}

# --- SIDEBAR (Unified) ---
with st.sidebar:
    st.title("🥣 specialk")
    if st.button("+ New Meeting", use_container_width=True):
        st.session_state.page = "new_meeting"
        st.rerun()

    # 1. Custom Prompts Tab
    with st.expander("⚡ Custom Prompts"):
        tab_add, tab_edit = st.tabs(["Add New", "Edit/Delete"])
        with tab_add:
            c1, c2 = st.columns([0.3, 0.7])
            new_icon = c1.text_input("Icon", value="💡", max_chars=2, key="add_ic")
            new_title = c2.text_input("Title", placeholder="Key Points", key="add_tl")
            new_text = st.text_area("Prompt Text", height=100, key="add_tx")
            if st.button("Add Prompt", key="add_p_btn", use_container_width=True):
                if new_title and new_text:
                    add_prompt(new_title, new_text, new_icon)
                    st.success("Added!"); time.sleep(0.5); st.rerun()
        with tab_edit:
            all_p = load_prompts()
            if not all_p: st.caption("No prompts found.")
            else:
                p_titles = [f"{p.get('icon','')} {p['title']}" for p in all_p]
                sel_p_idx = st.selectbox("Select Prompt", range(len(all_p)), format_func=lambda x: p_titles[x])
                if sel_p_idx is not None:
                    target = all_p[sel_p_idx]
                    with st.form(key=f"edit_form_{sel_p_idx}"):
                        ed_icon = st.text_input("Icon", value=target.get('icon', '⚡'), max_chars=2)
                        ed_title = st.text_input("Title", value=target['title'])
                        ed_text = st.text_area("Text", value=target['text'], height=100)
                        c_save, c_del = st.columns(2)
                        if c_save.form_submit_button("💾 Update"):
                            update_prompt(sel_p_idx, ed_title, ed_text, ed_icon)
                            st.success("Updated!"); time.sleep(0.5); st.rerun()
                        if c_del.form_submit_button("🗑️ Delete", type="primary"):
                            delete_prompt(sel_p_idx)
                            st.warning("Deleted!"); time.sleep(0.5); st.rerun()


    # 3. Settings & History
    with st.expander("⚙️ System Settings"):
        current = load_settings()
        saved_llms = current.get("saved_llm_models", ["local-model"])
        active_llm = current.get("llm_model_id", "local-model")

        col_in, col_add = st.columns([0.7, 0.3])
        new_id_input = col_in.text_input("New Model ID", placeholder="e.g. llama3", label_visibility="collapsed")
        if col_add.button("Add", key="global_add_llm"):
            if new_id_input and new_id_input not in saved_llms:
                saved_llms.append(new_id_input); current["saved_llm_models"] = saved_llms
                current["llm_model_id"] = new_id_input; save_settings(current); st.rerun()

        try: d_idx = saved_llms.index(active_llm)
        except ValueError: d_idx = 0
        st.selectbox("Active LLM", saved_llms, index=d_idx, key="active_llm_widget", on_change=update_model_setting)

        st.divider()
        st.subheader("🎤 Audio Transcription")

        # Engine selection - MLX Whisper as default (best quality)
        engs = ["MLX Whisper (Apple Silicon)", "WhisperX (Torch)"]
        curr_eng = current.get("transcription_engine", "MLX Whisper (Apple Silicon)")
        sel_eng = st.radio("Engine", engs, index=engs.index(curr_eng) if curr_eng in engs else 0)

        # Model selection based on engine
        if "WhisperX" in sel_eng:
            model_options = WHISPERX_MODELS
            default_model = "large-v3"
        else:
            model_options = MLX_MODELS
            default_model = "mlx-community/whisper-large-v3-turbo"

        curr_model = current.get("transcription_model", default_model)
        model_names = list(model_options.keys())
        model_values = list(model_options.values())

        # Find current model index
        try:
            model_idx = model_values.index(curr_model)
        except ValueError:
            model_idx = 0

        sel_model_name = st.selectbox("Model", model_names, index=model_idx)
        sel_model = model_options[sel_model_name]

        if st.button("💾 Save Audio Config"):
            current = load_settings()
            current["transcription_engine"] = sel_eng
            current["transcription_model"] = sel_model
            save_settings(current)
            st.success("Saved! Restart app to apply changes.")
            time.sleep(0.5)
            st.rerun()

        st.divider()
        st.subheader("🤖 LLM Transcription Formatter")
        st.caption("Use an LLM to clean and format raw transcriptions into clear, readable text.")

        llm_trans_enabled = current.get("llm_transcription_enabled", False)
        llm_trans_model = current.get("llm_transcription_model_id", "local-model")

        new_llm_trans_enabled = st.toggle(
            "Enable LLM Transcription",
            value=llm_trans_enabled,
            help="When enabled, an LLM will periodically format the raw transcription into clean, readable text."
        )

        # LLM Transcription Model selector
        try:
            llm_trans_idx = saved_llms.index(llm_trans_model)
        except ValueError:
            llm_trans_idx = 0

        new_llm_trans_model = st.selectbox(
            "LLM Transcription Model",
            saved_llms,
            index=llm_trans_idx,
            help="Select the LLM model to use for formatting transcriptions. Can be different from the chat model."
        )

        st.caption("ℹ️ LLM formats transcription every 10 seconds with 2-minute context window.")

        if st.button("💾 Save LLM Transcription Config"):
            current = load_settings()
            current["llm_transcription_enabled"] = new_llm_trans_enabled
            current["llm_transcription_model_id"] = new_llm_trans_model
            save_settings(current)
            st.success("LLM Transcription settings saved!")
            time.sleep(0.5)
            st.rerun()

        # Visible warnings if libs are missing
        if pypdf is None: st.error("❌ pypdf missing")
        if docx is None: st.error("❌ python-docx missing")

    st.divider()
    st.caption("History")
    meetings = load_meetings()

    if not meetings:
        st.caption("No meetings yet.")
    else:
        for m in meetings:
            col_btn, col_del = st.columns([0.85, 0.15])
            with col_btn:
                if st.button(f"{m['title']}\n{m['date']}", key=f"m_{m['id']}", use_container_width=True):
                    st.session_state.current_meeting_id = m['id']
                    st.session_state.chat_history = m.get("chat_history", [])
                    st.session_state.page = "review"
                    st.rerun()
            with col_del:
                if st.button("🗑️", key=f"del_{m['id']}", help=f"Delete {m['title']}"):
                    delete_meeting(m['id'])
                    st.rerun()

# --- PAGE: NEW MEETING ---
if st.session_state.page == "new_meeting":
    st.title("🎙️ Start New Meeting")
    with st.form("setup_form"):
        title = st.text_input("Meeting Title", value="Untitled Meeting")
        devices, default_idx = recorder.get_devices()
        if not devices: st.error("No microphones found."); sel_idx = None
        else:
            dev_names = [d[1] for d in devices]
            list_idx = 0
            for i, dev in enumerate(devices):
                if dev[0] == default_idx: list_idx = i; break
            sel_name = st.selectbox("Microphone", dev_names, index=list_idx)
            for dev in devices:
                if dev[1] == sel_name: sel_idx = dev[0]; break
        up_files = st.file_uploader("Upload Context", accept_multiple_files=True)
        if st.form_submit_button("Start Recording"):
            if sel_idx is None: st.error("Select mic.")
            else:
                recorder.start(sel_idx); st.session_state.temp_title = title; st.session_state.temp_files = up_files; st.session_state.chat_history = []; st.session_state.page = "recording"; st.rerun()

# --- PAGE: RECORDING ---
elif st.session_state.page == "recording":
    # Initialize live_up in session state if not exists
    if "live_uploaded_files" not in st.session_state:
        st.session_state.live_uploaded_files = []

    # Reset LLM transcription manager on new recording session
    if "recording_session_started" not in st.session_state:
        st.session_state.recording_session_started = True
        llm_transcription_manager.reset()

    c1, c2 = st.columns([0.8, 0.2])
    c1.title(f"🔴 Live: {st.session_state.temp_title}")
    if c2.button("🛑 Stop & Save", type="primary"):
        recorder.stop()
        # Parse context one last time to save it fully
        all_files = st.session_state.temp_files + st.session_state.live_uploaded_files
        ctx_str, _ = get_text_from_files(all_files)

        # Get the final transcript (use LLM formatted if available and enabled)
        current_settings = load_settings()
        llm_trans_enabled = current_settings.get("llm_transcription_enabled", False)
        final_transcript = recorder.get_transcript_text()
        if llm_trans_enabled:
            llm_formatted = llm_transcription_manager.get_formatted_output()
            if llm_formatted:
                final_transcript = f"=== LLM FORMATTED TRANSCRIPT ===\n{llm_formatted}\n\n=== RAW TRANSCRIPT ===\n{recorder.get_transcript_text()}"

        rec = save_meeting(st.session_state.temp_title, final_transcript, all_files, st.session_state.chat_history)
        st.session_state.live_uploaded_files = []  # Clear live files
        st.session_state.recording_session_started = False  # Reset session flag
        st.session_state.current_meeting_id = rec['id']; st.session_state.page = "review"; st.rerun()

    # Check if LLM transcription is enabled
    current_settings = load_settings()
    llm_trans_enabled = current_settings.get("llm_transcription_enabled", False)
    llm_trans_model = current_settings.get("llm_transcription_model_id", "local-model")

    if llm_trans_enabled:
        # Dual-box layout: LLM formatted (primary/larger) + Raw whisper (secondary/smaller)
        c_llm, c_raw, c_chat = st.columns([0.35, 0.25, 0.4])

        with c_llm:
            # Header with font size slider
            col_title, col_slider = st.columns([0.6, 0.4])
            col_title.subheader("📝 Formatted Transcript")

            # Font size slider - store in session state
            if "llm_font_size" not in st.session_state:
                st.session_state.llm_font_size = 16
            font_size = col_slider.slider("Font", 12, 24, st.session_state.llm_font_size, key="font_slider", label_visibility="collapsed")
            st.session_state.llm_font_size = font_size

            @st.fragment(run_every=2)
            def show_llm_formatted():
                if recorder.running:
                    raw_text = recorder.get_transcript_text()

                    # Get current output FIRST (never blocks)
                    current_output = llm_transcription_manager.get_formatted_output()

                    # Trigger formatting in background if needed (non-blocking check)
                    if not llm_transcription_manager.is_formatting:
                        llm_transcription_manager.update_and_format(raw_text)

                    # Get font size from session state
                    fs = st.session_state.get("llm_font_size", 16)
                    ts_fs = max(11, fs - 3)  # Timestamp slightly smaller

                    # Always display current output (or waiting message if empty)
                    if current_output:
                        # Clean display - render paragraphs nicely
                        paragraphs = current_output.strip().split('\n\n')
                        html_parts = []

                        for para in paragraphs:
                            if not para.strip():
                                continue
                            escaped = html.escape(para.strip()).replace('\n', '<br>')

                            # Style timestamp if present at start
                            if escaped.startswith('[') and ']' in escaped:
                                ts_end = escaped.index(']') + 1
                                timestamp = escaped[:ts_end]
                                content = escaped[ts_end:].strip()
                                if content.startswith('<br>'):
                                    content = content[4:]
                                html_parts.append(
                                    f'<p style="margin:0 0 20px 0;line-height:1.8;">'
                                    f'<span style="color:#89b4fa;font-size:{ts_fs}px;font-weight:500;">{timestamp}</span> '
                                    f'{content}'
                                    f'</p>'
                                )
                            else:
                                html_parts.append(
                                    f'<p style="margin:0 0 20px 0;line-height:1.8;">{escaped}</p>'
                                )

                        content_html = ''.join(html_parts) if html_parts else f'<p>{html.escape(current_output)}</p>'

                        # Auto-scroll JavaScript - use unique ID to force execution on each render
                        scroll_id = f"scroll_{int(time.time() * 1000)}"
                        scroll_script = f"""
                        <script id="{scroll_id}">
                            (function() {{
                                var container = document.getElementById('llm-transcript-box');
                                if (container) {{
                                    container.scrollTop = container.scrollHeight;
                                }}
                            }})();
                        </script>
                        """

                        st.markdown(
                            f'<div id="llm-transcript-box" style="height:500px;overflow-y:auto;background:#1e1e2e;padding:20px 24px;border-radius:12px;color:#cdd6f4;font-size:{fs}px;">{content_html}</div>{scroll_script}',
                            unsafe_allow_html=True
                        )
                    else:
                        st.markdown(
                            '<div style="height:500px;overflow-y:auto;background:#1e1e2e;padding:15px;border-radius:12px;display:flex;align-items:center;justify-content:center;color:#6c7086;">'
                            '<div style="text-align:center;">'
                            '<div style="font-size:32px;margin-bottom:10px;">🎙️</div>'
                            '<em>Waiting for transcription...</em><br>'
                            '<small>LLM will format after ~40 words</small>'
                            '</div></div>',
                            unsafe_allow_html=True
                        )
            show_llm_formatted()

        with c_raw:
            # Show active audio engine and model
            active_engine = current_settings.get("transcription_engine", "MLX Whisper (Apple Silicon)")
            active_model = current_settings.get("transcription_model", "mlx-community/whisper-large-v3-turbo")
            model_display = active_model.split("/")[-1] if "/" in active_model else active_model

            st.subheader("🎤 Raw Transcript")
            st.caption(f"**Engine:** {active_engine.split('(')[0].strip()}")

            @st.fragment(run_every=1)
            def show_raw_live():
                if recorder.running:
                    txt = html.escape(recorder.get_transcript_text()).replace("\n", "<br>")
                    st.markdown(
                        f'<div style="height:500px;overflow-y:auto;background:#333;padding:10px;border-radius:8px;font-size:12px;opacity:0.9;">{txt}</div>',
                        unsafe_allow_html=True
                    )
            show_raw_live()

        with c_chat:
            st.subheader("💬 Chat")
            u_in = st.chat_input("Ask...")

            # --- PREPARE DEBUG INFO ---
            ctx_transcript = recorder.get_transcript_text()
            all_context_files = st.session_state.temp_files + st.session_state.live_uploaded_files
            ctx_files, errs = get_text_from_files(all_context_files)

            # Store in session state for debug display
            debug_text = ""
            if ctx_transcript.strip():
                debug_text += f"=== MEETING TRANSCRIPT (Primary Content) ===\n{ctx_transcript}"
            if ctx_files.strip():
                debug_text += f"\n\n=== REFERENCE DOCUMENTS (Supporting Context) ===\n{ctx_files}"
            st.session_state.debug_context_text = debug_text
            st.session_state.debug_errors = errs

            for msg in st.session_state.chat_history: st.chat_message(msg["role"]).write(msg["content"])

            if u_in:
                st.session_state.chat_history.append({"role": "user", "content": u_in})
                st.chat_message("user").write(u_in)
                with st.chat_message("assistant"):
                    resp = ask_llm(u_in, ctx_transcript, ctx_files, use_history=False)
                    st.write(resp)
                st.session_state.chat_history.append({"role": "assistant", "content": resp}); st.rerun()

            # Add Live Files - under the chat
            with st.expander("📎 Add Context Files", expanded=False):
                live_up = st.file_uploader("Upload files to add context", accept_multiple_files=True, key="live_files_uploader")
                if live_up:
                    existing_names = [f.name for f in st.session_state.live_uploaded_files]
                    for f in live_up:
                        if f.name not in existing_names:
                            st.session_state.live_uploaded_files.append(f)

                if st.session_state.temp_files or st.session_state.live_uploaded_files:
                    st.caption("**Loaded files:**")
                    for f in st.session_state.temp_files:
                        st.write(f"📄 {f.name} *(initial)*")
                    for f in st.session_state.live_uploaded_files:
                        st.write(f"📄 {f.name} *(live)*")

            # Debug Context Window (under chat)
            with st.expander("🛠️ Debug Context", expanded=False):
                st.caption("This shows exactly what the LLM sees.")
                if "debug_context_text" in st.session_state:
                    word_count = len(st.session_state.debug_context_text.split())
                    st.caption(f"📊 **Total words in context: {word_count:,}**")

                    if "debug_system_prompt" in st.session_state:
                        st.text_area("System Prompt:", st.session_state.debug_system_prompt, height=300, disabled=True)

                    if "debug_user_prompt" in st.session_state:
                        st.text_area("User Prompt:", st.session_state.debug_user_prompt, height=100, disabled=True)

                    if "debug_errors" in st.session_state and st.session_state.debug_errors:
                        for e in st.session_state.debug_errors: st.error(e)
                else:
                    st.info("No context loaded yet. Send a message to see the full prompt.")

    else:
        # Original layout when LLM transcription is disabled
        c_trans, c_chat = st.columns([0.5, 0.5])
        with c_trans:
            # Show active audio engine and model
            active_engine = current_settings.get("transcription_engine", "MLX Whisper (Apple Silicon)")
            active_model = current_settings.get("transcription_model", "mlx-community/whisper-large-v3-turbo")
            model_display = active_model.split("/")[-1] if "/" in active_model else active_model

            st.subheader("📝 Transcript")
            st.caption(f"🎤 **Engine:** {active_engine} | **Model:** `{model_display}`")

            @st.fragment(run_every=1)
            def show_live():
                if recorder.running:
                    txt = html.escape(recorder.get_transcript_text()).replace("\n", "<br>")
                    st.markdown(f'<div style="height:550px;overflow-y:auto;background:#222;padding:10px;border-radius:8px;">{txt}</div>', unsafe_allow_html=True)
            show_live()

        with c_chat:
            st.subheader("💬 Chat")
            u_in = st.chat_input("Ask...")

            # --- PREPARE DEBUG INFO ---
            ctx_transcript = recorder.get_transcript_text()
            all_context_files = st.session_state.temp_files + st.session_state.live_uploaded_files
            ctx_files, errs = get_text_from_files(all_context_files)

            # Store in session state for debug display
            debug_text = ""
            if ctx_transcript.strip():
                debug_text += f"=== MEETING TRANSCRIPT (Primary Content) ===\n{ctx_transcript}"
            if ctx_files.strip():
                debug_text += f"\n\n=== REFERENCE DOCUMENTS (Supporting Context) ===\n{ctx_files}"
            st.session_state.debug_context_text = debug_text
            st.session_state.debug_errors = errs

            for msg in st.session_state.chat_history: st.chat_message(msg["role"]).write(msg["content"])

            if u_in:
                st.session_state.chat_history.append({"role": "user", "content": u_in})
                st.chat_message("user").write(u_in)
                with st.chat_message("assistant"):
                    resp = ask_llm(u_in, ctx_transcript, ctx_files, use_history=False)
                    st.write(resp)
                st.session_state.chat_history.append({"role": "assistant", "content": resp}); st.rerun()

            # Add Live Files - under the chat
            with st.expander("📎 Add Context Files", expanded=False):
                live_up = st.file_uploader("Upload files to add context", accept_multiple_files=True, key="live_files_uploader_std")
                if live_up:
                    existing_names = [f.name for f in st.session_state.live_uploaded_files]
                    for f in live_up:
                        if f.name not in existing_names:
                            st.session_state.live_uploaded_files.append(f)

                if st.session_state.temp_files or st.session_state.live_uploaded_files:
                    st.caption("**Loaded files:**")
                    for f in st.session_state.temp_files:
                        st.write(f"📄 {f.name} *(initial)*")
                    for f in st.session_state.live_uploaded_files:
                        st.write(f"📄 {f.name} *(live)*")

            # Debug Context Window (under chat)
            with st.expander("🛠️ Debug Context", expanded=False):
                st.caption("This shows exactly what the LLM sees.")
                if "debug_context_text" in st.session_state:
                    word_count = len(st.session_state.debug_context_text.split())
                    st.caption(f"📊 **Total words in context: {word_count:,}**")

                    if "debug_system_prompt" in st.session_state:
                        st.text_area("System Prompt:", st.session_state.debug_system_prompt, height=300, disabled=True)

                    if "debug_user_prompt" in st.session_state:
                        st.text_area("User Prompt:", st.session_state.debug_user_prompt, height=100, disabled=True)

                    if "debug_errors" in st.session_state and st.session_state.debug_errors:
                        for e in st.session_state.debug_errors: st.error(e)
                else:
                    st.info("No context loaded yet. Send a message to see the full prompt.")

# --- PAGE: REVIEW ---
elif st.session_state.page == "review":
    ms = load_meetings()
    m = next((x for x in ms if x['id'] == st.session_state.current_meeting_id), None)

    if not m: st.error("Meeting lost."); st.button("Home", on_click=lambda: setattr(st.session_state, 'page', 'home'))
    else:
        st.title(f"🥣 {m['title']}")
        t_chat, t_trans, t_file = st.tabs(["Chat", "Transcript", "Files"])

        with t_trans:
            if os.path.exists(m['transcript_path']): st.text_area("Transcript", open(m['transcript_path']).read(), height=600)

        with t_file:
            if m.get('context_files'):
                for p in m['context_files']:
                    if os.path.exists(p): st.write(f"📄 {os.path.basename(p)}")
            st.divider()
            new_up = st.file_uploader("Add Context", accept_multiple_files=True)
            if st.button("Update"):
                if new_up:
                    add_files_to_meeting(m['id'], new_up)
                    st.success("Files added!"); time.sleep(0.5); st.rerun()

        with t_chat:
            active_llm_now = load_settings().get("llm_model_id", "local-model")
            st.subheader(f"Chat (Using: `{active_llm_now}`)")

            with st.popover("⚙️ Style"):
                st.session_state.style_detail = st.selectbox("Detail", list(STYLE_PROMPTS["detail"].keys()))
                st.session_state.style_tone = st.selectbox("Tone", list(STYLE_PROMPTS["tone"].keys()))
                st.session_state.style_lang = st.selectbox("Language", list(STYLE_PROMPTS["language"].keys()))

            prompts = get_top_prompts(10)
            cols = st.columns(5); p_sel = None
            for i, p in enumerate(prompts):
                if cols[i % 5].button(f"{p.get('icon','⚡')} {p['title']}", key=f"rev_p_{i}"): p_sel = p['text']

            # --- PREPARE CONTEXT & DEBUG ---
            transcript_ctx = ""
            files_ctx = ""
            errors = []

            # 1. Transcript
            if os.path.exists(m['transcript_path']):
                transcript_ctx = open(m['transcript_path']).read()

            # 2. Files
            if m.get('context_files'):
                for fp in m['context_files']:
                    if os.path.exists(fp):
                        fname = os.path.basename(fp).lower()
                        try:
                            f_text = ""
                            if fname.endswith('.pdf'):
                                if pypdf:
                                    with open(fp, "rb") as f:
                                        pdf = pypdf.PdfReader(f)
                                        for p in pdf.pages:
                                            extracted = p.extract_text()
                                            if extracted: f_text += extracted + "\n"
                                else: errors.append(f"Missing pypdf for {fname}")
                            elif fname.endswith('.docx'):
                                if docx:
                                    doc = docx.Document(fp)
                                    f_text = "\n".join([p.text for p in doc.paragraphs])
                                else: errors.append(f"Missing python-docx for {fname}")
                            elif fname.endswith(('.txt', '.md', '.csv', '.py')):
                                with open(fp, "r", encoding="utf-8", errors="ignore") as f: f_text = f.read()

                            if f_text: files_ctx += f"\n\n--- FILE: {fname} ---\n{f_text}"
                            else:
                                if fname.endswith(('.pdf', '.docx', '.txt')):
                                    errors.append(f"⚠️ {fname} is empty (Scanned?).")
                        except Exception as e: errors.append(f"Error {fname}: {e}")

            # Build debug text with clear separation
            debug_text = ""
            if transcript_ctx.strip():
                debug_text += f"=== MEETING TRANSCRIPT (Primary Content) ===\n{transcript_ctx}"
            if files_ctx.strip():
                debug_text += f"\n\n=== REFERENCE DOCUMENTS (Supporting Context) ===\n{files_ctx}"
            st.session_state.debug_context_text = debug_text
            st.session_state.debug_errors = errors

            # Chat UI
            chat_cont = st.container(height=450)
            with chat_cont:
                for msg in st.session_state.chat_history:
                    with st.chat_message(msg["role"]): st.markdown(msg["content"])

            u_in = st.chat_input("Ask...")
            # Use prompt button if clicked
            f_query = p_sel if p_sel else u_in

            if f_query:
                # Increment usage count if it's a custom prompt
                if p_sel:
                    increment_prompt_usage(p_sel)

                st.session_state.chat_history.append({"role": "user", "content": f_query})
                with chat_cont:
                    with st.chat_message("user"): st.write(f_query)

                with chat_cont:
                    with st.chat_message("assistant"):
                        with st.spinner("Thinking..."):
                            resp = ask_llm(f_query, transcript_ctx, files_ctx, use_history=False); st.markdown(resp)
                st.session_state.chat_history.append({"role": "assistant", "content": resp}); st.rerun()

            # Debug Context Window (under chat)
            with st.expander("🛠️ Debug Context", expanded=False):
                st.caption("This shows exactly what the LLM sees.")
                if "debug_context_text" in st.session_state:
                    word_count = len(st.session_state.debug_context_text.split())
                    st.caption(f"📊 **Total words in context: {word_count:,}**")

                    # Show system prompt if available
                    if "debug_system_prompt" in st.session_state:
                        st.text_area("System Prompt:", st.session_state.debug_system_prompt, height=300, disabled=True)

                    # Show user prompt if available
                    if "debug_user_prompt" in st.session_state:
                        st.text_area("User Prompt:", st.session_state.debug_user_prompt, height=100, disabled=True)

                    if "debug_errors" in st.session_state and st.session_state.debug_errors:
                        for e in st.session_state.debug_errors: st.error(e)
                else:
                    st.info("No context loaded yet. Send a message to see the full prompt.")

elif st.session_state.page == "home":
    st.header("Welcome to specialk"); st.markdown("Use the sidebar to start a meeting.")
