"""
Special K - Parallel Intelligence Meeting Assistant

A synchronized split-view interface with:
- Fast Track (Left Pane): Live Transcription
- Slow Track (Right Pane): Deep Analysis (Topics, Decisions, Actions)

Single Brain Architecture: One shared AI model with priority queuing.
"""

import streamlit as st
import time
import requests
import os
import html
import threading
import json
import re
from datetime import datetime

try:
    import pypdf
except ImportError:
    pypdf = None
try:
    import docx
except ImportError:
    docx = None

from audio_engine import AudioRecorder, MLX_MODELS
from meeting_manager import (
    load_meetings, save_meeting, delete_meeting, update_meeting_chat,
    load_prompts, add_prompt, delete_prompt, update_prompt,
    load_settings, save_settings, add_files_to_meeting,
    increment_prompt_usage, get_top_prompts, update_meeting_title
)
from analysis_engine import (
    AnalysisEngine, InsightType, TopicState
)
from llm_provider import (
    get_provider, test_provider, ClaudeProvider, LocalLLMProvider
)

# --- CONFIG ---
st.set_page_config(page_title="specialk", layout="wide", page_icon="🥣")
LLM_URL = "http://localhost:1234/v1/chat/completions"
CLAUDE_MODELS = list(ClaudeProvider.MODELS.keys())

# --- CUSTOM CSS FOR FULL-HEIGHT CONTAINERS ---
st.markdown("""
<style>
/* Make main containers fill viewport height */
.full-height-container {
    height: calc(100vh - 200px);
    overflow-y: auto;
    padding: 8px;
}

.full-height-container-chat {
    height: calc(100vh - 350px);
    overflow-y: auto;
}

/* Waiting state container */
.full-height-waiting {
    height: calc(100vh - 200px);
    background: #1e1e2e;
    padding: 20px;
    border-radius: 12px;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #6c7086;
}
</style>
""", unsafe_allow_html=True)

# --- HELPER: AUTO-SAVE CALLBACK ---
def update_model_setting():
    if "active_llm_widget" in st.session_state:
        new_model = st.session_state.active_llm_widget
        current_settings = load_settings()
        current_settings["llm_model_id"] = new_model
        save_settings(current_settings)

# --- HELPER: FILE PARSER ---
# Reads text from uploaded files to build LLM context
def get_text_from_files(file_list):
    context_str = ""
    errors = []

    for file in file_list:
        fname = file.name.lower()
        try:
            file.seek(0)
            text = ""

            # 1. PDF Handling
            if fname.endswith('.pdf'):
                if pypdf:
                    pdf = pypdf.PdfReader(file)
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted: text += extracted + "\n"
                    context_str += f"\n\n--- FILE: {file.name} (PDF) ---\n{text}"
                else:
                    errors.append(f"❌ Missing 'pypdf' library. Cannot read {file.name}")

            # 2. Word Doc Handling
            elif fname.endswith('.docx'):
                if docx:
                    doc = docx.Document(file)
                    text = "\n".join([p.text for p in doc.paragraphs])
                    context_str += f"\n\n--- FILE: {file.name} (DOCX) ---\n{text}"
                else:
                    errors.append(f"❌ Missing 'python-docx' library. Cannot read {file.name}")

            # 3. Plain Text Handling
            elif fname.endswith(('.txt', '.md', '.csv', '.json', '.py')):
                text = file.read().decode("utf-8")
                context_str += f"\n\n--- FILE: {file.name} ---\n{text}"

            # Validation
            if not text and not errors and fname.endswith(('.pdf', '.docx', '.txt')):
                errors.append(f"⚠️ {file.name} was empty or unreadable (Scanned?).")

        except Exception as e:
            errors.append(f"❌ Error reading {file.name}: {str(e)}")

    return context_str, errors

# --- HELPER: LLM FUNCTION ---
def ask_llm(prompt, transcript_text, files_text, use_history=False):
    settings = load_settings()
    llm_id = settings.get("llm_model_id", "local-model")
    all_prompts = load_prompts()
    is_custom = any(p['text'].strip() == prompt.strip() for p in all_prompts)

    # Load style preferences (with defaults)
    if "style_detail" not in st.session_state: st.session_state.style_detail = "Brief & Concise (Default)"
    if "style_tone" not in st.session_state: st.session_state.style_tone = "Neutral (Default)"
    if "style_lang" not in st.session_state: st.session_state.style_lang = "Human (Default)"

    # Build context with clear separation
    context_block = ""
    if transcript_text.strip():
        context_block += f"\n\n=== MEETING TRANSCRIPT (Primary Content) ===\n{transcript_text}"
    if files_text.strip():
        context_block += f"\n\n=== REFERENCE DOCUMENTS (Supporting Context) ===\n{files_text}"

    base_instructions = """You are a meeting assistant.
IMPORTANT: The MEETING TRANSCRIPT is the primary content - this is what the meeting is about.
The REFERENCE DOCUMENTS are supplementary materials uploaded for context/reference only.
When asked questions, focus on the MEETING TRANSCRIPT not the referenced documents.
Use the REFERENCE DOCUMENTS only to provide additional context, provide better answers or clarify topics discussed in the meeting."""

    if is_custom:
        system_content = f"{base_instructions}{context_block}"
    else:
        detail_instr = STYLE_PROMPTS["detail"][st.session_state.style_detail]
        tone_instr = STYLE_PROMPTS["tone"][st.session_state.style_tone]
        lang_instr = STYLE_PROMPTS["language"][st.session_state.style_lang]
        style_block = f"{detail_instr} {tone_instr} {lang_instr}".strip()
        system_content = f"{base_instructions}\nSTYLE: {style_block}{context_block}"

    # Store the full prompt for debug display
    st.session_state.debug_system_prompt = system_content
    st.session_state.debug_user_prompt = prompt

    messages = [{"role": "system", "content": system_content}]
    if use_history:
        messages.extend(st.session_state.chat_history)
    messages.append({"role": "user", "content": prompt})

    try:
        provider = get_provider(settings)
        start_time = time.perf_counter()
        content = provider.chat_with_messages(messages, temperature=0.7, timeout=300)
        elapsed = round(time.perf_counter() - start_time, 2)

        if content:
            label = f"{provider.provider_name()}: {provider.model_name()}"
            return f"{content}\n\n---\n*⏱️ {elapsed}s | {label}*"
        else: return "Error: Empty response from LLM"
    except Exception as e: return f"Connection Error: {e}"


# --- HELPER: STRICT OUTPUT CLEANER ---
def _extract_clean_transcript(llm_output):
    """
    Aggressively extract only the actual transcript content from LLM output.
    Removes ALL meta-commentary, preambles, speaker labels, and formatting artifacts.
    """
    if not llm_output:
        return ""

    result = llm_output

    # 1. Remove any content before the first actual sentence
    # LLMs often start with "Here's..." or "The cleaned transcript:"
    preamble_patterns = [
        r'^.*?(?:here\'s|here is|the cleaned|the formatted|below is|i\'ve cleaned|cleaned version)[^:]*:\s*',
        r'^.*?(?:transcript|text|output)[^:]*:\s*',
    ]
    for pattern in preamble_patterns:
        result = re.sub(pattern, '', result, flags=re.IGNORECASE | re.DOTALL)

    # 2. Remove speaker labels - we have NO speaker data, these are hallucinations
    # Patterns: "[00:01:02] You:", "[00:01:02] Me:", "[00:01:02] Speaker:", "Speaker 1:", etc.
    speaker_patterns = [
        r'\[[\d:]+\]\s*(?:You|Me|Speaker\s*\d*|Person\s*\d*|Host|Guest|Interviewer|Interviewee)\s*:\s*',
        r'^(?:You|Me|Speaker\s*\d*|Person\s*\d*|Host|Guest|Interviewer|Interviewee)\s*:\s*',
        r'\n(?:You|Me|Speaker\s*\d*|Person\s*\d*|Host|Guest|Interviewer|Interviewee)\s*:\s*',
    ]
    for pattern in speaker_patterns:
        result = re.sub(pattern, '\n', result, flags=re.IGNORECASE | re.MULTILINE)

    # 3. Remove bullet points that are meta-commentary
    lines = result.split('\n')
    cleaned_lines = []
    meta_patterns = [
        r'^\s*[\*\-•]\s*(?:added|changed|removed|fixed|broke|converted|replaced|combined|split|moved|deleted|corrected|standardized|preserved|removing|correcting|keeping|i kept|i made|note:|changes:)',
        r'^\s*(?:notes?:|changes?:|edit(?:s|ed)?:|summary:|explanation:)',
        r'^\s*(?:i hope|let me know|feel free|alternatively|if you|or,?\s*if)',
    ]

    for line in lines:
        line_lower = line.lower().strip()
        should_skip = False

        for pattern in meta_patterns:
            if re.match(pattern, line_lower):
                should_skip = True
                break

        # Skip lines that are just formatting markers
        if line.strip() in ['---', '***', '===', '```', '```text', '```transcript']:
            should_skip = True

        if not should_skip:
            cleaned_lines.append(line)

    result = '\n'.join(cleaned_lines)

    # 4. Clean up excessive whitespace
    result = re.sub(r'\n{3,}', '\n\n', result)
    result = result.strip()

    return result


def _is_valid_transcript_output(text, original_input_length):
    """
    Validate that LLM output looks like cleaned transcript, not hallucination.
    Returns (is_valid, reason)
    """
    if not text or len(text) < 20:
        return False, "Output too short"

    # Check for excessive length (hallucination indicator)
    if len(text) > original_input_length * 1.5:
        return False, f"Output too long ({len(text)} vs input {original_input_length})"

    # Check for common hallucination patterns
    hallucination_indicators = [
        r'\[[\d:]+\]\s*(?:You|Me|Speaker)',  # Hallucinated speaker labels
        r'(?:You|Me|Speaker\s*\d*):\s',  # Speaker attribution
        r'(?:here\'s|here is) the (?:cleaned|formatted)',  # Preamble leaked through
    ]

    for pattern in hallucination_indicators:
        if re.search(pattern, text, re.IGNORECASE):
            return False, f"Contains hallucination pattern: {pattern}"

    return True, "OK"


class LLMTranscriptionManager:
    """
    Simplified, robust LLM transcription formatter.

    Key design principles:
    1. SIMPLE: Single-pass cleaning, no complex context revision
    2. STRICT: JSON output format to prevent preambles
    3. VALIDATED: Always validate output before accepting
    4. FALLBACK: Raw text is always better than bad LLM output

    The previous architecture tried to be too clever with "working paragraphs"
    that could be revised - this led to confusion and hallucinations.
    """

    def __init__(self):
        self.formatted_chunks = []  # Already formatted and validated chunks
        self.is_formatting = False
        self.lock = threading.Lock()
        self.last_raw_length_processed = 0
        self.min_chars_for_format = 300  # Process after ~300 chars (more context = better formatting)
        self.start_time = None
        self.failed_attempts = 0  # Track consecutive failures

    def reset(self):
        """Reset the manager for a new session."""
        with self.lock:
            self.formatted_chunks = []
            self.is_formatting = False
            self.last_raw_length_processed = 0
            self.start_time = time.time()
            self.failed_attempts = 0

    def _get_timestamp(self):
        """Get current timestamp based on elapsed recording time."""
        if self.start_time is None:
            self.start_time = time.time()
        elapsed = int(time.time() - self.start_time)
        hours = elapsed // 3600
        minutes = (elapsed % 3600) // 60
        seconds = elapsed % 60
        return f"[{hours:02d}:{minutes:02d}:{seconds:02d}]"

    def _clean_raw_text(self, text):
        """Remove cursor indicator and normalize whitespace."""
        if not text:
            return ""
        # Remove cursor and timestamps for clean processing
        text = text.replace(' ▌', '').replace('▌', '')
        return text.strip()

    def _strip_timestamps(self, text):
        """Remove all timestamps from text for LLM processing."""
        return re.sub(r'\[\d{2}:\d{2}:\d{2}\]\s*', '', text)

    def should_format(self, current_raw_text):
        """Check if we have enough new text to trigger formatting."""
        clean_raw = self._clean_raw_text(current_raw_text)
        new_chars = len(clean_raw) - self.last_raw_length_processed
        return new_chars >= self.min_chars_for_format and not self.is_formatting

    def _call_llm_for_cleaning(self, text_to_clean):
        """
        Call LLM with a strict, simple prompt focused ONLY on cleaning.
        No speaker attribution, no complex formatting, just grammar fixes.
        """
        settings = load_settings()

        system_prompt = """You are a transcript cleaner. Your ONLY job is to fix grammar and remove filler words.

RULES:
1. Fix grammar, punctuation, capitalization
2. Remove filler words: um, uh, like, you know, sort of, kind of, I mean, basically
3. Remove stutters and repeated words
4. Group 3-6 related sentences into paragraphs
5. Separate paragraphs with ONE blank line

CRITICAL - DO NOT:
- Add ANY words not in the original
- Add speaker labels (You:, Me:, Speaker:)
- Add introductions ("Here's the cleaned...")
- Add explanations or notes
- Repeat any content

OUTPUT: Just the cleaned text. Nothing else."""

        user_prompt = f"Clean this:\n\n{text_to_clean}"

        try:
            provider = get_provider(settings, purpose="transcription")
            return provider.chat(system_prompt, user_prompt, temperature=0.0, timeout=60)
        except Exception as e:
            print(f"LLM Error: {e}")
            return None

    def _do_format_in_background(self, raw_text):
        """Background thread: format new text with strict validation."""
        try:
            clean_raw = self._clean_raw_text(raw_text)

            if len(clean_raw) <= self.last_raw_length_processed:
                return

            # Get the new portion of text to process
            new_text = clean_raw[self.last_raw_length_processed:].strip()
            new_text_no_timestamps = self._strip_timestamps(new_text)

            if not new_text_no_timestamps or len(new_text_no_timestamps) < 50:
                return

            # Skip garbage/noise input - don't waste LLM calls on nonsense
            if self._is_noise_input(new_text_no_timestamps):
                print(f"⏭️ Skipping noise input: {new_text_no_timestamps[:50]}...")
                with self.lock:
                    self.last_raw_length_processed = len(clean_raw)
                return

            # Call LLM for cleaning
            llm_result = self._call_llm_for_cleaning(new_text_no_timestamps)

            if llm_result:
                # Extract clean transcript (remove any leaked meta-commentary)
                cleaned = _extract_clean_transcript(llm_result)

                # Validate the output
                is_valid, reason = _is_valid_transcript_output(cleaned, len(new_text_no_timestamps))

                if is_valid:
                    # Add timestamp and store
                    ts = self._get_timestamp()
                    timestamped_chunk = f"{ts} {cleaned}"

                    with self.lock:
                        self.formatted_chunks.append(timestamped_chunk)
                        self.last_raw_length_processed = len(clean_raw)
                        self.failed_attempts = 0

                    print(f"LLM formatting successful: {len(cleaned)} chars")
                else:
                    print(f"LLM output validation failed: {reason}")
                    self.failed_attempts += 1

                    # After 3 consecutive failures, use raw text as fallback
                    if self.failed_attempts >= 3:
                        print("Too many failures, using raw text as fallback")
                        ts = self._get_timestamp()
                        # Simple cleanup without LLM
                        simple_cleaned = self._simple_cleanup(new_text_no_timestamps)
                        timestamped_chunk = f"{ts} {simple_cleaned}"

                        with self.lock:
                            self.formatted_chunks.append(timestamped_chunk)
                            self.last_raw_length_processed = len(clean_raw)
                            self.failed_attempts = 0
            else:
                # LLM call failed
                self.failed_attempts += 1
                print(f"LLM call failed (attempt {self.failed_attempts})")

                # Fallback to raw text after failures
                if self.failed_attempts >= 3:
                    ts = self._get_timestamp()
                    simple_cleaned = self._simple_cleanup(new_text_no_timestamps)
                    timestamped_chunk = f"{ts} {simple_cleaned}"

                    with self.lock:
                        self.formatted_chunks.append(timestamped_chunk)
                        self.last_raw_length_processed = len(clean_raw)
                        self.failed_attempts = 0

        except Exception as e:
            print(f"LLM Transcription Manager Error: {e}")
        finally:
            with self.lock:
                self.is_formatting = False

    def _is_noise_input(self, text):
        """
        Detect if input is garbage/noise that shouldn't be sent to the LLM.
        Returns True if the text appears to be noise (e.g., repeated syllables, gibberish).
        """
        if not text:
            return True

        # Normalize
        text_lower = text.lower().strip()
        words = text_lower.split()

        if len(words) < 3:
            return True

        # Check for highly repetitive content (same word repeated many times)
        word_counts = {}
        for word in words:
            # Strip punctuation for comparison
            clean_word = ''.join(c for c in word if c.isalnum())
            if clean_word:
                word_counts[clean_word] = word_counts.get(clean_word, 0) + 1

        if word_counts:
            max_count = max(word_counts.values())
            # If one word appears more than 60% of the time, it's likely noise
            if max_count / len(words) > 0.6:
                return True

        # Check for very short repeated syllables (like "ep ep ep" or "ah ah ah")
        short_words = [w for w in words if len(w) <= 3]
        if len(short_words) / len(words) > 0.7 and len(words) > 5:
            return True

        # Check average word length - real speech has longer words
        avg_word_len = sum(len(w) for w in words) / len(words)
        if avg_word_len < 2.5 and len(words) > 5:
            return True

        return False

    def _simple_cleanup(self, text):
        """
        Simple non-LLM cleanup as fallback.
        Just removes obvious filler words and normalizes whitespace.
        """
        if not text:
            return ""

        # Remove common filler words
        fillers = [
            r'\b(um+|uh+|er+|ah+)\b',
            r'\b(you know)\b',
            r'\b(sort of|kind of)\b',
            r'\b(I mean)\b',
            r'\b(basically)\b',
        ]

        result = text
        for pattern in fillers:
            result = re.sub(pattern, '', result, flags=re.IGNORECASE)

        # Clean up multiple spaces and normalize
        result = re.sub(r'\s+', ' ', result)
        result = result.strip()

        return result

    def update_and_format(self, current_raw_text):
        """Non-blocking: triggers formatting in background if needed."""
        if not current_raw_text or len(current_raw_text.strip()) < 50:
            return

        if not self.should_format(current_raw_text):
            return

        with self.lock:
            if self.is_formatting:
                return
            self.is_formatting = True

        format_thread = threading.Thread(
            target=self._do_format_in_background,
            args=(current_raw_text,),
            daemon=True
        )
        format_thread.start()

    def get_formatted_output(self):
        """Get the current formatted output as a single string."""
        with self.lock:
            if not self.formatted_chunks:
                return ""
            return "\n\n".join(self.formatted_chunks)

    def get_last_processed_position(self):
        """Get the position in raw text that has been processed by LLM."""
        with self.lock:
            return self.last_raw_length_processed


# --- STATE & SETTINGS ---
@st.cache_resource
def get_recorder(): return AudioRecorder(load_settings())
recorder = get_recorder()

@st.cache_resource
def get_llm_transcription_manager():
    return LLMTranscriptionManager()
llm_transcription_manager = get_llm_transcription_manager()

@st.cache_resource
def get_analysis_engine():
    settings = load_settings()
    llm_id = settings.get("llm_model_id", "local-model")
    provider = get_provider(settings, purpose="chat")
    return AnalysisEngine(LLM_URL, llm_id, provider=provider)
analysis_engine = get_analysis_engine()

if "page" not in st.session_state: st.session_state.page = "home"
if "current_meeting_id" not in st.session_state: st.session_state.current_meeting_id = None
if "chat_history" not in st.session_state: st.session_state.chat_history = []
if "show_all_prompts" not in st.session_state: st.session_state.show_all_prompts = False
if "temp_title" not in st.session_state: st.session_state.temp_title = "Untitled"
if "temp_files" not in st.session_state: st.session_state.temp_files = []
if "live_uploaded_files" not in st.session_state: st.session_state.live_uploaded_files = []
if "selected_segment_id" not in st.session_state: st.session_state.selected_segment_id = None
if "seek_to_segment" not in st.session_state: st.session_state.seek_to_segment = None

STYLE_PROMPTS = {
    "detail": {
        "Brief & Concise (Default)": "Provide short, focused responses. Use bullet points when listing multiple items. Limit explanations to 2-3 sentences maximum. Omit unnecessary context or background information.",
        "Brief (One Sentence)": "Respond with exactly one sentence. Be direct and to the point. Do not elaborate or provide additional context.",
        "Detailed": "Provide comprehensive, thorough responses. Include relevant context, examples, and explanations. Break down complex topics into clear sections. Feel free to elaborate on important points."
    },
    "tone": {
        "Neutral (Default)": "Use a professional, balanced tone. Be objective and straightforward without being cold or overly formal.",
        "Friendly": "Use a warm, conversational tone. Be approachable and personable while remaining helpful. Use casual language where appropriate.",
        "Formal": "Use a formal, professional tone suitable for business communication. Avoid contractions and colloquialisms. Maintain a respectful, polished style."
    },
    "language": {
        "Human (Default)": "Write naturally as a human would. Use varied sentence structures, everyday vocabulary, and a conversational flow. Avoid robotic or overly structured phrasing.",
        "AI Normal": "You may use structured formatting, technical terminology, and systematic organization. Clarity and precision take priority over conversational style."
    }
}

# --- SIDEBAR (Unified) ---
with st.sidebar:
    st.title("🥣 specialk")
    if st.button("+ New Meeting", use_container_width=True):
        st.session_state.page = "new_meeting"
        st.rerun()

    # 1. Custom Prompts Tab
    with st.expander("⚡ Custom Prompts"):
        tab_add, tab_edit = st.tabs(["Add New", "Edit/Delete"])
        with tab_add:
            c1, c2 = st.columns([0.3, 0.7])
            new_icon = c1.text_input("Icon", value="💡", max_chars=2, key="add_ic")
            new_title = c2.text_input("Title", placeholder="Key Points", key="add_tl")
            new_text = st.text_area("Prompt Text", height=100, key="add_tx")
            if st.button("Add Prompt", key="add_p_btn", use_container_width=True):
                if new_title and new_text:
                    add_prompt(new_title, new_text, new_icon)
                    st.success("Added!"); time.sleep(0.5); st.rerun()
        with tab_edit:
            all_p = load_prompts()
            if not all_p: st.caption("No prompts found.")
            else:
                p_titles = [f"{p.get('icon','')} {p['title']}" for p in all_p]
                sel_p_idx = st.selectbox("Select Prompt", range(len(all_p)), format_func=lambda x: p_titles[x])
                if sel_p_idx is not None:
                    target = all_p[sel_p_idx]
                    with st.form(key=f"edit_form_{sel_p_idx}"):
                        ed_icon = st.text_input("Icon", value=target.get('icon', '⚡'), max_chars=2)
                        ed_title = st.text_input("Title", value=target['title'])
                        ed_text = st.text_area("Text", value=target['text'], height=100)
                        c_save, c_del = st.columns(2)
                        if c_save.form_submit_button("💾 Update"):
                            update_prompt(sel_p_idx, ed_title, ed_text, ed_icon)
                            st.success("Updated!"); time.sleep(0.5); st.rerun()
                        if c_del.form_submit_button("🗑️ Delete", type="primary"):
                            delete_prompt(sel_p_idx)
                            st.warning("Deleted!"); time.sleep(0.5); st.rerun()


    # 3. Settings & History
    with st.expander("⚙️ System Settings"):
        current = load_settings()
        saved_llms = current.get("saved_llm_models", ["local-model"])
        active_llm = current.get("llm_model_id", "local-model")

        col_in, col_add = st.columns([0.7, 0.3])
        new_id_input = col_in.text_input("New Model ID", placeholder="e.g. llama3", label_visibility="collapsed")
        if col_add.button("Add", key="global_add_llm"):
            if new_id_input and new_id_input not in saved_llms:
                saved_llms.append(new_id_input); current["saved_llm_models"] = saved_llms
                current["llm_model_id"] = new_id_input; save_settings(current); st.rerun()

        try: d_idx = saved_llms.index(active_llm)
        except ValueError: d_idx = 0

        col_select, col_del = st.columns([0.8, 0.2])
        col_select.selectbox("Active LLM", saved_llms, index=d_idx, key="active_llm_widget", on_change=update_model_setting)

        # Delete model button
        if col_del.button("🗑️", key="delete_llm_btn", help="Delete selected model"):
            selected_model = st.session_state.get("active_llm_widget")
            if selected_model and len(saved_llms) > 1:
                saved_llms.remove(selected_model)
                current["saved_llm_models"] = saved_llms
                # If deleted model was active, switch to first available
                if active_llm == selected_model:
                    current["llm_model_id"] = saved_llms[0]
                # Also update transcription model if it was using deleted model
                if current.get("llm_transcription_model_id") == selected_model:
                    current["llm_transcription_model_id"] = saved_llms[0]
                save_settings(current)
                st.rerun()
            elif len(saved_llms) <= 1:
                st.warning("Cannot delete the last model!")

        st.divider()
        st.subheader("☁️ LLM Provider")
        st.caption("Choose between a local LLM or Claude API for chat & analysis.")

        provider_options = {"Local LLM": "local", "Claude API": "claude"}
        current_provider = current.get("llm_provider", "local")
        provider_labels = list(provider_options.keys())
        provider_values = list(provider_options.values())
        try:
            prov_idx = provider_values.index(current_provider)
        except ValueError:
            prov_idx = 0

        sel_provider_label = st.selectbox("Chat Provider", provider_labels, index=prov_idx)
        sel_provider = provider_options[sel_provider_label]

        if sel_provider == "claude":
            claude_key = current.get("claude_api_key", "")
            new_claude_key = st.text_input(
                "Anthropic API Key",
                value=claude_key,
                type="password",
                help="Get your key from console.anthropic.com"
            )

            claude_model_id = current.get("claude_model_id", "claude-sonnet-4-20250514")
            claude_model_labels = {v: k for k, v in ClaudeProvider.MODELS.items()}
            claude_display = list(ClaudeProvider.MODELS.values())
            try:
                cm_idx = list(ClaudeProvider.MODELS.keys()).index(claude_model_id)
            except ValueError:
                cm_idx = 0
            sel_claude_display = st.selectbox("Claude Model", claude_display, index=cm_idx)
            sel_claude_id = claude_model_labels[sel_claude_display]

            if st.button("💾 Save Provider Config"):
                current = load_settings()
                current["llm_provider"] = sel_provider
                current["claude_api_key"] = new_claude_key
                current["claude_model_id"] = sel_claude_id
                save_settings(current)
                st.success("Claude provider saved!")
                time.sleep(0.5)
                st.rerun()

            if st.button("🧪 Test Claude Connection"):
                test_settings = {**current, "llm_provider": "claude", "claude_api_key": new_claude_key, "claude_model_id": sel_claude_id}
                test_prov = get_provider(test_settings)
                ok, msg = test_provider(test_prov)
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)
        else:
            if st.button("💾 Save Provider Config"):
                current = load_settings()
                current["llm_provider"] = "local"
                save_settings(current)
                st.success("Local LLM provider saved!")
                time.sleep(0.5)
                st.rerun()

        # Transcription provider (can be different from chat)
        st.divider()
        st.subheader("🎤 Audio Transcription")

        model_options = MLX_MODELS
        default_model = "mlx-community/whisper-large-v3-turbo"

        curr_model = current.get("transcription_model", default_model)
        model_names = list(model_options.keys())
        model_values = list(model_options.values())

        try:
            model_idx = model_values.index(curr_model)
        except ValueError:
            model_idx = 0

        sel_model_name = st.selectbox("Model", model_names, index=model_idx)
        sel_model = model_options[sel_model_name]

        if st.button("💾 Save Audio Config"):
            current = load_settings()
            current["transcription_model"] = sel_model
            save_settings(current)
            st.success("Saved! Restart app to apply changes.")
            time.sleep(0.5)
            st.rerun()

        st.divider()
        st.subheader("🤖 LLM Transcription Formatter")
        st.caption("Use an LLM to clean and format raw transcriptions into clear, readable text.")

        trans_provider_options = {"Local LLM": "local", "Claude API": "claude"}
        current_trans_provider = current.get("llm_transcription_provider", "local")
        try:
            tp_idx = list(trans_provider_options.values()).index(current_trans_provider)
        except ValueError:
            tp_idx = 0

        sel_trans_provider_label = st.selectbox("Transcription Provider", list(trans_provider_options.keys()), index=tp_idx)
        sel_trans_provider = trans_provider_options[sel_trans_provider_label]

        llm_trans_enabled = current.get("llm_transcription_enabled", False)

        new_llm_trans_enabled = st.toggle(
            "Enable LLM Transcription",
            value=llm_trans_enabled,
            help="When enabled, an LLM will periodically format the raw transcription into clean, readable text."
        )

        if sel_trans_provider == "local":
            llm_trans_model = current.get("llm_transcription_model_id", "local-model")
            try:
                llm_trans_idx = saved_llms.index(llm_trans_model)
            except ValueError:
                llm_trans_idx = 0

            new_llm_trans_model = st.selectbox(
                "LLM Transcription Model",
                saved_llms,
                index=llm_trans_idx,
                help="Select the local LLM model to use for formatting transcriptions."
            )
        else:
            claude_trans_model_id = current.get("claude_transcription_model_id", "claude-sonnet-4-20250514")
            claude_trans_display = list(ClaudeProvider.MODELS.values())
            claude_trans_labels = {v: k for k, v in ClaudeProvider.MODELS.items()}
            try:
                ctm_idx = list(ClaudeProvider.MODELS.keys()).index(claude_trans_model_id)
            except ValueError:
                ctm_idx = 0
            sel_ctrans_display = st.selectbox("Claude Transcription Model", claude_trans_display, index=ctm_idx)
            new_llm_trans_model = claude_trans_labels[sel_ctrans_display]

        st.caption("ℹ️ LLM formats transcription every 10 seconds with 2-minute context window.")

        if st.button("💾 Save LLM Transcription Config"):
            current = load_settings()
            current["llm_transcription_enabled"] = new_llm_trans_enabled
            current["llm_transcription_provider"] = sel_trans_provider
            if sel_trans_provider == "local":
                current["llm_transcription_model_id"] = new_llm_trans_model
            else:
                current["claude_transcription_model_id"] = new_llm_trans_model
            save_settings(current)
            st.success("LLM Transcription settings saved!")
            time.sleep(0.5)
            st.rerun()

        # Visible warnings if libs are missing
        if pypdf is None: st.error("❌ pypdf missing")
        if docx is None: st.error("❌ python-docx missing")

    st.divider()
    st.caption("History")
    meetings = load_meetings()

    if not meetings:
        st.caption("No meetings yet.")
    else:
        for m in meetings:
            col_btn, col_del = st.columns([0.85, 0.15])
            with col_btn:
                if st.button(f"{m['title']}\n{m['date']}", key=f"m_{m['id']}", use_container_width=True):
                    st.session_state.current_meeting_id = m['id']
                    st.session_state.chat_history = m.get("chat_history", [])
                    st.session_state.page = "review"
                    st.rerun()
            with col_del:
                if st.button("🗑️", key=f"del_{m['id']}", help=f"Delete {m['title']}"):
                    delete_meeting(m['id'])
                    st.rerun()

# --- PAGE: NEW MEETING ---
if st.session_state.page == "new_meeting":
    st.title("🎙️ Start New Meeting")
    with st.form("setup_form"):
        title = st.text_input("Meeting Title", value="Untitled Meeting")
        devices, default_idx = recorder.get_devices()
        if not devices: st.error("No microphones found."); sel_idx = None
        else:
            dev_names = [d[1] for d in devices]
            list_idx = 0
            for i, dev in enumerate(devices):
                if dev[0] == default_idx: list_idx = i; break
            sel_name = st.selectbox("Microphone", dev_names, index=list_idx)
            for dev in devices:
                if dev[1] == sel_name: sel_idx = dev[0]; break
        up_files = st.file_uploader("Upload Context", accept_multiple_files=True)
        if st.form_submit_button("Start Recording"):
            if sel_idx is None: st.error("Select mic.")
            else:
                # Reset and start the analysis engine
                analysis_engine.reset()
                analysis_engine.start()
                llm_transcription_manager.reset()
                recorder.start(sel_idx)
                st.session_state.temp_title = title
                st.session_state.temp_files = up_files
                st.session_state.chat_history = []
                st.session_state.selected_segment_id = None
                st.session_state.recording_session_started = True
                st.session_state.page = "recording"
                st.rerun()

# --- PAGE: RECORDING (PARALLEL INTELLIGENCE MODE) ---
elif st.session_state.page == "recording":
    # Initialize session state
    if "live_uploaded_files" not in st.session_state:
        st.session_state.live_uploaded_files = []

    # Reset managers on new recording session
    if "recording_session_started" not in st.session_state:
        st.session_state.recording_session_started = True
        analysis_engine.reset()
        analysis_engine.start()
        llm_transcription_manager.reset()

    # Header row
    c1, c2, c3 = st.columns([0.6, 0.2, 0.2])

    with c1:
        new_title = st.text_input("Meeting Title", value=st.session_state.temp_title, label_visibility="collapsed", key="live_meeting_title")
        if new_title != st.session_state.temp_title:
            st.session_state.temp_title = new_title

    c2.markdown("### 🔴 Recording")

    if c3.button("🛑 Stop & Save", type="primary"):
        recorder.stop()
        analysis_engine.stop()

        all_files = st.session_state.temp_files + st.session_state.live_uploaded_files
        ctx_str, _ = get_text_from_files(all_files)

        # Get raw transcript from recorder
        raw_transcript = recorder.committed_text if hasattr(recorder, 'committed_text') else ""

        # Get LLM-formatted transcript from the transcription manager
        llm_transcript = llm_transcription_manager.get_formatted_output()

        # If LLM transcription hasn't processed yet, use raw as fallback
        if not llm_transcript or not llm_transcript.strip():
            llm_transcript = raw_transcript

        # Get analysis state
        state = analysis_engine.get_state()

        # Build analysis data JSON
        analysis_data = json.dumps({
            "topics": state["topics"],
            "insights": state["insights"]
        }, indent=2)

        # Save with both transcripts and analysis
        full_content = f"""=== LLM TRANSCRIPT ===
{llm_transcript}

=== RAW TRANSCRIPT ===
{raw_transcript}

=== ANALYSIS DATA ===
{analysis_data}
"""
        rec = save_meeting(st.session_state.temp_title, full_content, all_files, st.session_state.chat_history)
        st.session_state.live_uploaded_files = []
        st.session_state.recording_session_started = False
        st.session_state.selected_segment_id = None
        st.session_state.current_meeting_id = rec['id']
        st.session_state.page = "review"
        st.rerun()

    st.markdown("---")

    # === SPLIT VIEW: LEFT (Transcript) | RIGHT (Analysis) ===
    c_left, c_right = st.columns([0.55, 0.45])

    # --- LEFT PANE: LIVE TRANSCRIPT ---
    with c_left:
        st.subheader("📝 Live Transcript")

        current_settings = load_settings()
        active_engine = current_settings.get("transcription_engine", "MLX Whisper (Apple Silicon)")

        # Engine info and font size slider
        col_engine, col_font = st.columns([0.6, 0.4])
        with col_engine:
            st.caption(f"🎤 {active_engine.split('(')[0].strip()}")
        with col_font:
            font_size = st.slider("Font", min_value=12, max_value=24, value=16, step=1, key="transcript_font_size", label_visibility="collapsed")

        @st.fragment(run_every=1)
        def show_live_transcript():
            fs = st.session_state.get("transcript_font_size", 16)
            ts_fs = max(10, fs - 4)  # Timestamp font slightly smaller

            if recorder.running:
                raw_text = recorder.get_transcript_text()

                # Trigger LLM formatting in background (non-blocking)
                if raw_text:
                    llm_transcription_manager.update_and_format(raw_text)

                # Only add FINALIZED segments to analysis engine (not live/growing text)
                committed_text = recorder.committed_text if hasattr(recorder, 'committed_text') else ""

                if committed_text:
                    # Parse committed (finalized) segments for analysis
                    committed_lines = committed_text.split('\n')
                    timestamp_pattern = r'\[(\d{2}:\d{2}:\d{2})\]'

                    for line in committed_lines:
                        if line.strip():
                            ts_match = re.match(timestamp_pattern, line)
                            timestamp = ts_match.group(1) if ts_match else None
                            clean_text = re.sub(timestamp_pattern, '', line).strip()

                            if clean_text and len(clean_text) > 20:
                                existing_segments = analysis_engine.transcript_manager.get_all_segments()
                                already_added = any(
                                    s.text.strip() == clean_text.strip()
                                    for s in existing_segments
                                )
                                if not already_added:
                                    analysis_engine.add_transcript_segment(clean_text, timestamp)

                # Get LLM-formatted transcript
                llm_formatted = llm_transcription_manager.get_formatted_output()

                # Get currently speaking text (text after the last timestamp that has cursor)
                live_speaking_text = ""
                if raw_text and '▌' in raw_text:
                    # Find the last timestamp and get text after it
                    cursor_pos = raw_text.rfind('▌')
                    text_before_cursor = raw_text[:cursor_pos].strip()
                    all_matches = list(re.finditer(r'\[(\d{2}:\d{2}:\d{2})\]', text_before_cursor))

                    if all_matches:
                        last_match = all_matches[-1]
                        live_speaking_text = text_before_cursor[last_match.end():].strip()
                    else:
                        live_speaking_text = text_before_cursor

                # Build display - ONLY LLM formatted + LIVE, never raw
                html_parts = []

                # Check if we need to seek to a segment
                seek_segment_id = st.session_state.get("seek_to_segment")
                seek_text = ""
                if seek_segment_id:
                    segment = analysis_engine.get_segment_by_id(seek_segment_id)
                    if segment:
                        seek_text = segment.text[:50].lower()  # First 50 chars for matching
                    # Clear the seek after processing
                    st.session_state.seek_to_segment = None

                # 1. Display LLM-formatted paragraphs ONLY (never raw text)
                if llm_formatted and llm_formatted.strip():
                    paragraphs = llm_formatted.strip().split('\n\n')
                    for idx, para in enumerate(paragraphs):
                        para = para.strip()
                        if not para:
                            continue

                        # Extract timestamp from paragraph if present
                        ts_match = re.match(r'\[(\d{2}:\d{2}:\d{2})\]\s*(.*)', para, re.DOTALL)
                        if ts_match:
                            timestamp = f"[{ts_match.group(1)}]"
                            text = ts_match.group(2).strip()
                        else:
                            timestamp = ""
                            text = para

                        if text:
                            escaped_text = html.escape(text)

                            # Check if this paragraph contains the seek text
                            is_highlighted = seek_text and seek_text in text.lower()

                            if is_highlighted:
                                # Highlighted style with green border
                                html_parts.append(f'''<div id="segment_{idx}" style="background:#2d4a2d;padding:14px 16px;margin-bottom:8px;border-radius:8px;border:2px solid #a6e3a1;"><span style="color:#89b4fa;font-size:{ts_fs}px;font-weight:600;">{timestamp}</span> <span style="color:#cdd6f4;font-size:{fs}px;line-height:1.7;">{escaped_text}</span></div>''')
                            elif timestamp:
                                html_parts.append(f'''<div style="background:#1e1e2e;padding:14px 16px;margin-bottom:8px;border-radius:8px;"><span style="color:#89b4fa;font-size:{ts_fs}px;font-weight:600;">{timestamp}</span> <span style="color:#cdd6f4;font-size:{fs}px;line-height:1.7;">{escaped_text}</span></div>''')
                            else:
                                html_parts.append(f'''<div style="background:#1e1e2e;padding:14px 16px;margin-bottom:8px;border-radius:8px;"><span style="color:#cdd6f4;font-size:{fs}px;line-height:1.7;">{escaped_text}</span></div>''')

                # 2. Display currently speaking text with LIVE badge (always show if speaking)
                if live_speaking_text:
                    # Check if live text overlaps with end of LLM formatted - remove overlap
                    if llm_formatted:
                        # Get last ~100 chars of LLM output for comparison
                        llm_ending = re.sub(r'\[\d{2}:\d{2}:\d{2}\]\s*', '', llm_formatted[-200:]).lower().strip()
                        live_lower = live_speaking_text.lower().strip()

                        # If live text is contained in LLM ending, don't show it
                        if live_lower and len(live_lower) > 10 and live_lower in llm_ending:
                            live_speaking_text = ""
                        # If first words of live match end of LLM, trim the overlap
                        elif live_lower:
                            live_words = live_lower.split()
                            for trim_count in range(min(10, len(live_words)), 0, -1):
                                overlap_check = ' '.join(live_words[:trim_count])
                                if overlap_check in llm_ending:
                                    # Trim overlapping words from live text
                                    original_words = live_speaking_text.split()
                                    live_speaking_text = ' '.join(original_words[trim_count:]).strip()
                                    break

                if live_speaking_text:
                    escaped_live = html.escape(live_speaking_text)
                    html_parts.append(f'''<div style="background:#3d2a1a;padding:14px 16px;margin-bottom:8px;border-radius:8px;border:1px solid #f5a623;"><span style="background:#f5a623;color:#1a1a1a;font-size:11px;font-weight:700;padding:2px 8px;border-radius:4px;margin-right:8px;">LIVE</span><span style="color:#cdd6f4;font-size:{fs}px;line-height:1.7;">{escaped_live}</span></div>''')

                if html_parts:
                    content = f'''<div class="full-height-container">{''.join(html_parts)}</div>'''
                    st.markdown(content, unsafe_allow_html=True)
                else:
                    st.markdown('''<div class="full-height-waiting"><div style="text-align:center;"><div style="font-size:48px;margin-bottom:15px;">🎙️</div><em>Waiting for transcription...</em></div></div>''', unsafe_allow_html=True)

        show_live_transcript()

    # --- RIGHT PANE: SLOW TRACK (DEEP ANALYSIS) ---
    with c_right:
        st.subheader("🧠 Slow Track - Deep Analysis")

        tab_topics, tab_decisions, tab_actions, tab_chat = st.tabs([
            "📌 Topics", "✅ Decisions", "📋 Actions", "💬 Chat"
        ])

        # --- TOPICS TAB ---
        with tab_topics:
            @st.fragment(run_every=3)
            def show_topics():
                state = analysis_engine.get_state()
                active_topics = state["topics"]["active"]
                resolved_topics = state["topics"]["resolved"]
                parking_lot = state["topics"]["parking_lot"]

                has_any_topics = active_topics or resolved_topics or parking_lot

                with st.container(height=400):
                    if active_topics:
                        st.markdown("##### 🟢 Active Topics")
                        for i, topic in enumerate(active_topics):
                            col1, col2 = st.columns([0.85, 0.15])
                            with col1:
                                st.markdown(f"**{topic['name']}**")
                                st.caption(f"Mentioned {topic['mention_count']}x")
                            with col2:
                                source_ids = topic.get("source_segment_ids", [])
                                if source_ids:
                                    if st.button("🔗", key=f"seek_active_{i}_{topic['id']}", help="Jump to source"):
                                        st.session_state.seek_to_segment = source_ids[0]
                            st.markdown("---")

                    if resolved_topics:
                        st.markdown("##### ✅ Resolved Topics")
                        for i, topic in enumerate(resolved_topics):
                            col1, col2 = st.columns([0.85, 0.15])
                            with col1:
                                st.markdown(f"~~{topic['name']}~~")
                                if topic.get("resolution_summary"):
                                    st.success(f"**Resolution:** {topic['resolution_summary']}")
                                else:
                                    st.caption("_Resolution details not captured_")
                            with col2:
                                source_ids = topic.get("source_segment_ids", [])
                                if source_ids:
                                    if st.button("🔗", key=f"seek_resolved_{i}_{topic['id']}", help="Jump to source"):
                                        st.session_state.seek_to_segment = source_ids[0]
                            st.markdown("---")

                    if parking_lot:
                        st.markdown("##### 🅿️ Parking Lot")
                        st.caption("Topics mentioned but not fully discussed")
                        for i, topic in enumerate(parking_lot):
                            col1, col2 = st.columns([0.85, 0.15])
                            with col1:
                                st.markdown(f"*{topic['name']}*")
                            with col2:
                                source_ids = topic.get("source_segment_ids", [])
                                if source_ids:
                                    if st.button("🔗", key=f"seek_parked_{i}_{topic['id']}", help="Jump to source"):
                                        st.session_state.seek_to_segment = source_ids[0]

                    if not has_any_topics:
                        st.caption("_Topics will appear here as they are detected in the conversation._")

            show_topics()

        # --- DECISIONS TAB ---
        with tab_decisions:
            @st.fragment(run_every=3)
            def show_decisions():
                state = analysis_engine.get_state()
                decisions = state["insights"]["decisions"]

                with st.container(height=400):
                    if decisions:
                        for i, dec in enumerate(decisions):
                            col1, col2 = st.columns([0.85, 0.15])
                            with col1:
                                st.markdown(f"✅ **{dec['content']}**")
                                st.caption(f"📍 {dec['timestamp']}")
                            with col2:
                                source_ids = dec.get("source_segment_ids", [])
                                if source_ids:
                                    if st.button("🔗", key=f"seek_dec_{i}_{dec['id']}", help="Jump to source"):
                                        st.session_state.seek_to_segment = source_ids[0]
                            st.markdown("---")
                    else:
                        st.caption("_Decisions will appear here as they are detected in the conversation._")

            show_decisions()

        # --- ACTIONS TAB ---
        with tab_actions:
            @st.fragment(run_every=3)
            def show_actions():
                state = analysis_engine.get_state()
                actions = state["insights"]["actions"]

                with st.container(height=400):
                    if actions:
                        for i, action in enumerate(actions):
                            col1, col2 = st.columns([0.85, 0.15])
                            with col1:
                                st.markdown(f"📋 **{action['content']}**")
                                st.caption(f"📍 {action['timestamp']}")
                            with col2:
                                source_ids = action.get("source_segment_ids", [])
                                if source_ids:
                                    if st.button("🔗", key=f"seek_action_{i}_{action['id']}", help="Jump to source"):
                                        st.session_state.seek_to_segment = source_ids[0]
                            st.markdown("---")
                    else:
                        st.caption("_Action items will appear here as they are detected in the conversation._")

            show_actions()

        # --- CHAT TAB (Fast Track priority) ---
        with tab_chat:
            st.caption("Ask questions about the meeting (Fast Track priority)")

            prompts = get_top_prompts(5)
            p_sel = None
            if prompts:
                cols = st.columns(min(5, len(prompts)))
                for i, p in enumerate(prompts[:5]):
                    if cols[i].button(f"{p.get('icon', '⚡')} {p['title']}", key=f"rec_p_{i}"):
                        p_sel = p['text']

            # Chat messages in scrollable container
            chat_html = ""
            for msg in st.session_state.chat_history:
                role_icon = "👤" if msg["role"] == "user" else "🤖"
                chat_html += f'<div style="margin-bottom:10px;padding:10px;background:{"#f0f0f0" if msg["role"] == "user" else "#ffffff"};border-radius:8px;border:1px solid #e0e0e0;color:#333;"><strong>{role_icon}</strong> {html.escape(msg["content"])}</div>'

            if chat_html:
                st.markdown(f'<div class="full-height-container-chat">{chat_html}</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="full-height-container-chat" style="display:flex;align-items:center;justify-content:center;color:#6c7086;"><em>No messages yet. Ask a question about the meeting.</em></div>', unsafe_allow_html=True)

            u_in = st.chat_input("Ask about the meeting...")
            f_query = p_sel if p_sel else u_in

            # Build context from analysis engine
            state = analysis_engine.get_state()
            ctx_transcript = state["transcript"]["full_text"]
            all_context_files = st.session_state.temp_files + st.session_state.live_uploaded_files
            ctx_files, errs = get_text_from_files(all_context_files)

            if f_query:
                if p_sel:
                    increment_prompt_usage(p_sel)
                st.session_state.chat_history.append({"role": "user", "content": f_query})
                with st.spinner("Thinking..."):
                    resp = ask_llm(f_query, ctx_transcript, ctx_files, use_history=False)
                st.session_state.chat_history.append({"role": "assistant", "content": resp})
                st.rerun()

    # --- CONTEXT FILES EXPANDER ---
    with st.expander("📎 Add Context Files", expanded=False):
        live_up = st.file_uploader("Upload files", accept_multiple_files=True, key="live_files_uploader")
        if live_up:
            existing_names = [f.name for f in st.session_state.live_uploaded_files]
            for f in live_up:
                if f.name not in existing_names:
                    st.session_state.live_uploaded_files.append(f)

        if st.session_state.temp_files or st.session_state.live_uploaded_files:
            st.caption("**Loaded files:**")
            for f in st.session_state.temp_files:
                st.write(f"📄 {f.name} *(initial)*")
            for f in st.session_state.live_uploaded_files:
                st.write(f"📄 {f.name} *(live)*")

# --- PAGE: REVIEW ---
elif st.session_state.page == "review":
    ms = load_meetings()
    m = next((x for x in ms if x['id'] == st.session_state.current_meeting_id), None)

    if not m: st.error("Meeting lost."); st.button("Home", on_click=lambda: setattr(st.session_state, 'page', 'home'))
    else:
        # Editable meeting title
        col_title, col_save = st.columns([0.85, 0.15])
        with col_title:
            new_title = st.text_input("Meeting Title", value=m['title'], label_visibility="collapsed", key="review_meeting_title")
        with col_save:
            if st.button("💾 Save Title", key="save_title_btn"):
                if new_title and new_title != m['title']:
                    update_meeting_title(m['id'], new_title)
                    st.success("Title updated!")
                    time.sleep(0.3)
                    st.rerun()

        # Parse transcript content
        llm_transcript = ""
        raw_transcript = ""
        analysis_data = None

        if os.path.exists(m['transcript_path']):
            full_content = open(m['transcript_path']).read()

            # New format: LLM TRANSCRIPT + RAW TRANSCRIPT + ANALYSIS DATA
            if "=== LLM TRANSCRIPT ===" in full_content:
                parts = full_content.split("=== RAW TRANSCRIPT ===")
                llm_transcript = parts[0].replace("=== LLM TRANSCRIPT ===", "").strip()

                if len(parts) > 1:
                    remaining = parts[1]
                    if "=== ANALYSIS DATA ===" in remaining:
                        raw_parts = remaining.split("=== ANALYSIS DATA ===")
                        raw_transcript = raw_parts[0].strip()
                        try:
                            analysis_data = json.loads(raw_parts[1].strip())
                        except json.JSONDecodeError:
                            analysis_data = None
                    else:
                        raw_transcript = remaining.strip()

            # Legacy format: TRANSCRIPT + ANALYSIS DATA
            elif "=== ANALYSIS DATA ===" in full_content:
                parts = full_content.split("=== ANALYSIS DATA ===")
                llm_transcript = parts[0].replace("=== TRANSCRIPT ===", "").strip()
                raw_transcript = llm_transcript  # Same as LLM for legacy
                try:
                    analysis_data = json.loads(parts[1].strip())
                except json.JSONDecodeError:
                    analysis_data = None

            # Older legacy format: LLM FORMATTED TRANSCRIPT + RAW TRANSCRIPT
            elif "=== LLM FORMATTED TRANSCRIPT ===" in full_content:
                parts = full_content.split("=== RAW TRANSCRIPT ===")
                llm_transcript = parts[0].replace("=== LLM FORMATTED TRANSCRIPT ===", "").strip()
                raw_transcript = parts[1].strip() if len(parts) > 1 else llm_transcript

            # Fallback: plain transcript
            else:
                llm_transcript = full_content
                raw_transcript = full_content

        t_transcript, t_raw, t_analysis, t_chat, t_file = st.tabs(["📝 Transcript", "🔊 Raw Audio", "🧠 Analysis", "💬 Chat", "📁 Files"])

        with t_transcript:
            st.subheader("LLM Processed Transcript")
            st.caption("Cleaned and formatted by AI")
            st.text_area("LLM Transcript", llm_transcript, height=500, label_visibility="collapsed")

        with t_raw:
            st.subheader("Raw Live Transcript")
            st.caption("Original transcription from audio engine")
            st.text_area("Raw Transcript", raw_transcript, height=500, label_visibility="collapsed")

        with t_analysis:
            if analysis_data:
                col1, col2 = st.columns(2)

                with col1:
                    st.subheader("📌 Topics")

                    if analysis_data.get("topics", {}).get("active"):
                        st.markdown("**🟢 Active**")
                        for t in analysis_data["topics"]["active"]:
                            st.markdown(f"- {t['name']}")

                    if analysis_data.get("topics", {}).get("resolved"):
                        st.markdown("**✅ Resolved**")
                        for t in analysis_data["topics"]["resolved"]:
                            st.markdown(f"- ~~{t['name']}~~")
                            if t.get("resolution_summary"):
                                st.caption(f"  → {t['resolution_summary']}")

                    if analysis_data.get("topics", {}).get("parking_lot"):
                        st.markdown("**🅿️ Parking Lot**")
                        for t in analysis_data["topics"]["parking_lot"]:
                            st.markdown(f"- *{t['name']}*")

                with col2:
                    st.subheader("📋 Insights")

                    if analysis_data.get("insights", {}).get("decisions"):
                        st.markdown("**✅ Decisions**")
                        for d in analysis_data["insights"]["decisions"]:
                            st.markdown(f"- {d['content']}")

                    if analysis_data.get("insights", {}).get("actions"):
                        st.markdown("**📋 Action Items**")
                        for a in analysis_data["insights"]["actions"]:
                            st.markdown(f"- {a['content']}")

                    if analysis_data.get("insights", {}).get("key_points"):
                        st.markdown("**💡 Key Points**")
                        for kp in analysis_data["insights"]["key_points"]:
                            st.markdown(f"- {kp['content']}")
            else:
                st.info("No analysis data available for this meeting. Analysis is only generated during live recordings.")

        with t_file:
            if m.get('context_files'):
                for p in m['context_files']:
                    if os.path.exists(p): st.write(f"📄 {os.path.basename(p)}")
            st.divider()
            new_up = st.file_uploader("Add Context", accept_multiple_files=True)
            if st.button("Update"):
                if new_up:
                    add_files_to_meeting(m['id'], new_up)
                    st.success("Files added!"); time.sleep(0.5); st.rerun()

        with t_chat:
            _chat_settings = load_settings()
            _chat_provider = get_provider(_chat_settings)
            st.subheader(f"Chat (Using: `{_chat_provider.provider_name()} — {_chat_provider.model_name()}`)")

            with st.popover("⚙️ Style"):
                st.session_state.style_detail = st.selectbox("Detail", list(STYLE_PROMPTS["detail"].keys()))
                st.session_state.style_tone = st.selectbox("Tone", list(STYLE_PROMPTS["tone"].keys()))
                st.session_state.style_lang = st.selectbox("Language", list(STYLE_PROMPTS["language"].keys()))

            prompts = get_top_prompts(10)
            p_sel = None
            cols = st.columns(min(5, len(prompts)))
            for i, p in enumerate(prompts[:5]):
                if cols[i].button(f"{p.get('icon','⚡')} {p['title']}", key=f"rev_p_{i}"): p_sel = p['text']
            if len(prompts) > 5:
                cols2 = st.columns(min(5, len(prompts) - 5))
                for i, p in enumerate(prompts[5:10]):
                    if cols2[i].button(f"{p.get('icon','⚡')} {p['title']}", key=f"rev_p2_{i}"): p_sel = p['text']

            # --- PREPARE CONTEXT & DEBUG ---
            files_ctx = ""
            errors = []

            # Files context
            if m.get('context_files'):
                for fp in m['context_files']:
                    if os.path.exists(fp):
                        fname = os.path.basename(fp).lower()
                        try:
                            f_text = ""
                            if fname.endswith('.pdf'):
                                if pypdf:
                                    with open(fp, "rb") as f:
                                        pdf = pypdf.PdfReader(f)
                                        for page in pdf.pages:
                                            extracted = page.extract_text()
                                            if extracted: f_text += extracted + "\n"
                                else: errors.append(f"Missing pypdf for {fname}")
                            elif fname.endswith('.docx'):
                                if docx:
                                    doc_obj = docx.Document(fp)
                                    f_text = "\n".join([para.text for para in doc_obj.paragraphs])
                                else: errors.append(f"Missing python-docx for {fname}")
                            elif fname.endswith(('.txt', '.md', '.csv', '.py')):
                                with open(fp, "r", encoding="utf-8", errors="ignore") as f: f_text = f.read()

                            if f_text: files_ctx += f"\n\n--- FILE: {fname} ---\n{f_text}"
                            else:
                                if fname.endswith(('.pdf', '.docx', '.txt')):
                                    errors.append(f"⚠️ {fname} is empty (Scanned?).")
                        except Exception as e: errors.append(f"Error {fname}: {e}")

            # Chat UI - full browser height with HTML
            chat_html = ""
            for msg in st.session_state.chat_history:
                role_icon = "👤" if msg["role"] == "user" else "🤖"
                chat_html += f'<div style="margin-bottom:10px;padding:10px;background:{"#f0f0f0" if msg["role"] == "user" else "#ffffff"};border-radius:8px;border:1px solid #e0e0e0;color:#333;"><strong>{role_icon}</strong> {html.escape(msg["content"])}</div>'

            if chat_html:
                st.markdown(f'<div class="full-height-container-chat">{chat_html}</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div class="full-height-container-chat" style="display:flex;align-items:center;justify-content:center;color:#6c7086;"><em>No messages yet. Ask a question about the meeting.</em></div>', unsafe_allow_html=True)

            u_in = st.chat_input("Ask...")
            f_query = p_sel if p_sel else u_in

            if f_query:
                if p_sel:
                    increment_prompt_usage(p_sel)

                st.session_state.chat_history.append({"role": "user", "content": f_query})
                with st.spinner("Thinking..."):
                    resp = ask_llm(f_query, llm_transcript, files_ctx, use_history=False)
                st.session_state.chat_history.append({"role": "assistant", "content": resp}); st.rerun()

elif st.session_state.page == "home":
    st.header("🥣 Welcome to Special K")
    st.markdown("""
### Parallel Intelligence Meeting Assistant

**Special K** is a local, privacy-focused meeting assistant with a synchronized split-view interface.

#### 🎯 Key Features

**Fast Track (Left Pane)** - Live Transcription
- Real-time speech-to-text using MLX Whisper
- Every sentence gets a unique ID for tracking
- Click-to-seek from insights to source

**Slow Track (Right Pane)** - Deep Analysis
- Topics with lifecycle states (Active, Resolved, Parking Lot)
- Decisions and Action Items extraction
- Rolling context for smart deduplication

#### 🧠 Single Brain Architecture

One shared AI model handles both:
1. **Priority 1**: Chat and live transcription (never lags)
2. **Priority 2**: Background analysis (runs when idle)

#### 🚀 Getting Started

1. Click **+ New Meeting** in the sidebar
2. Select your microphone
3. Start recording and watch insights appear!
    """)
